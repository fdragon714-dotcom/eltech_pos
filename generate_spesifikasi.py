"""Generate Spesifikasi Tabel Database (3NF) Word document for POS El-Tech."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Data: tables from 3NF normalization
# Each entry: (group, table_no, table_title, table_name, columns, description)
# columns: list of (attribute, note)  where note in {"PK","FK",""}
# ---------------------------------------------------------------------------
DATA = [
    # ---------------- Entitas Master ----------------
    ("Entitas Master", "5.3", "Tabel User", "tb_user",
     [("id_user", "PK"), ("username", ""), ("email", ""), ("password", ""),
      ("last_login", ""), ("is_superuser", ""), ("is_staff", ""),
      ("is_active", ""), ("date_joined", "")],
     "Tabel user digunakan untuk menyimpan data yang berkaitan langsung dengan proses login dan "
     "otorisasi pengguna dalam sistem. Kolom username, email, password, serta atribut status akun "
     "seperti is_active, is_staff, dan is_superuser digunakan oleh sistem sebagai bagian dari "
     "mekanisme autentikasi dan pengelolaan hak akses. Tabel ini bersifat inti dan tidak menyimpan "
     "informasi personal pengguna secara lengkap, sehingga proses validasi sesi login dapat berjalan "
     "secara ringan, efisien, dan tetap aman."),

    ("Entitas Master", "5.4", "Tabel Kategori", "tb_kategori",
     [("id_kategori", "PK"), ("nama_kategori", ""), ("deskripsi_kategori", ""),
      ("gambar_kategori", ""), ("status_kategori", ""),
      ("tanggal_tambah_kategori", ""), ("tanggal_update_kategori", "")],
     "Tabel kategori digunakan untuk menyimpan data pengelompokan barang yang lebih spesifik di "
     "bawah klasifikasi tipe. Kolom-kolom seperti nama kategori, deskripsi, gambar, dan status "
     "keberadaan data berperan dalam menjaga keteraturan katalog produk dalam sistem Point of "
     "Sales, sehingga data tetap konsisten dan mudah dikelola."),

    ("Entitas Master", "5.5", "Tabel Tipe Kategori", "tb_tipe",
     [("id_tipe", "PK"), ("id_kategori", "FK"), ("nama_tipe", ""), ("gambar_tipe", "")],
     "Tabel tipe digunakan untuk menyimpan data jenis atau golongan barang secara umum sebagai "
     "klasifikasi tertinggi dalam hierarki inventaris. Tabel ini berfungsi sebagai data master yang "
     "menjadi acuan bagi tabel kategori di bawahnya melalui kolom id_kategori sebagai Foreign Key. "
     "Dengan memisahkan data tipe ke dalam tabel tersendiri, sistem terhindar dari penulisan "
     "klasifikasi yang berulang pada setiap entri barang, sehingga data tetap konsisten dan mudah "
     "dikelola saat dibutuhkan dalam proses pemfilteran maupun pelaporan."),

    ("Entitas Master", "5.6", "Tabel Produk", "tb_produk",
     [("id_produk", "PK"), ("id_kategori", "FK"), ("kode_produk", ""),
      ("nama_produk", ""), ("deskripsi_produk", ""), ("harga_jual", ""),
      ("hpp", ""), ("stok_produk", ""), ("garansi_hari", ""),
      ("status_produk", ""), ("gambar_produk", ""),
      ("tanggal_tambah_produk", ""), ("tanggal_update_produk", "")],
     "Tabel produk digunakan untuk menyimpan seluruh informasi barang yang tersedia dalam sistem, "
     "baik data fisik maupun data komersialnya. Kolom-kolom seperti kode produk, nama barang, harga "
     "jual, jumlah stok, dan informasi garansi (warranty) menjadi atribut utama yang diakses oleh "
     "sistem pada saat proses transaksi berlangsung. Tabel ini juga memiliki relasi ke tabel "
     "kategori melalui kunci asing id_kategori, sesuai dengan prinsip normalisasi basis data, guna "
     "menghindari redundansi data sekaligus mempercepat waktu eksekusi kueri ketika data produk "
     "dibutuhkan."),

    ("Entitas Master", "5.7", "Tabel Supplier", "tb_supplier",
     [("id_supplier", "PK"), ("nama_supplier", ""), ("kontak", ""),
      ("alamat", ""), ("status_supplier", ""),
      ("tanggal_tambah_supplier", ""), ("tanggal_update_supplier", "")],
     "Tabel supplier digunakan untuk menyimpan data pemasok barang. Tabel ini memuat informasi "
     "seperti nama supplier, kontak, dan alamat, yang berguna dalam proses pengadaan (pembelian) "
     "barang ke dalam inventaris sistem."),

    # ---------------- Entitas Penjualan ----------------
    ("Entitas Penjualan", "5.8", "Tabel Transaksi", "tb_transaksi",
     [("id_transaksi", "PK"), ("kode_transaksi", ""), ("sub_total", ""),
      ("grand_total", ""), ("diskon_amount", ""), ("diskon", ""),
      ("tendered_amount", ""), ("amount_change", ""), ("payment_method", ""),
      ("customer_name", ""), ("customer_wa", ""), ("customer_address", ""),
      ("id_user", "FK"), ("tanggal_tambah_transaksi", ""),
      ("tanggal_update_transaksi", "")],
     "Tabel transaksi digunakan untuk menyimpan data utama (header) dari setiap transaksi "
     "penjualan yang terjadi dalam sistem. Kolom-kolom seperti sub_total, grand_total, diskon, dan "
     "metode pembayaran merekam rekapitulasi finansial dari setiap transaksi, sedangkan kolom "
     "identitas konsumen mencatat pihak yang melakukan pembelian. Tabel ini juga memiliki relasi ke "
     "tabel user melalui id_user untuk mencatat kasir yang memvalidasi transaksi, sehingga setiap "
     "aktivitas transaksi dapat ditelusuri dan dipertanggungjawabkan secara akurat."),

    ("Entitas Penjualan", "5.9", "Tabel Transaksi Detail", "tb_transaksi_detail",
     [("id_item", "PK"), ("id_transaksi", "FK"), ("id_produk", "FK"),
      ("harga_item", ""), ("qty_beli", ""), ("total_item", "")],
     "Tabel transaksi detail digunakan untuk menyimpan rincian item barang dari setiap transaksi "
     "yang tercatat di tabel transaksi. Tabel ini berperan sebagai tabel perantara (bridge table) "
     "yang menangani relasi many-to-many antara transaksi dan produk, di mana setiap baris "
     "mewakili satu item barang beserta kolom qty_beli dan harga_item pada faktur tersebut. Nilai "
     "harga_item disimpan secara permanen pada saat transaksi dieksekusi, sehingga perubahan harga "
     "di kemudian hari tidak memengaruhi riwayat transaksi yang telah tercatat sebelumnya."),

    # ---------------- Entitas Pembelian (Inventory) ----------------
    ("Entitas Pembelian (Inventory)", "5.10", "Tabel Pembelian", "tb_pembelian",
     [("id_pembelian", "PK"), ("kode_pembelian", ""), ("id_supplier", "FK"),
      ("grand_total_pembelian", ""), ("tanggal_tambah_pembelian", ""),
      ("tanggal_update_pembelian", "")],
     "Tabel pembelian digunakan untuk mencatat data utama (header) setiap aktivitas restock atau "
     "pembelian barang dari supplier. Data total pembelian dan referensi supplier (id_supplier) "
     "disimpan pada tabel ini sebagai rekam jejak arus barang masuk."),

    ("Entitas Pembelian (Inventory)", "5.11", "Tabel Pembelian Detail", "tb_pembelian_detail",
     [("id_item_masuk", "PK"), ("id_pembelian", "FK"), ("id_produk", "FK"),
      ("harga_masuk", ""), ("qty_masuk", ""), ("total_masuk", "")],
     "Tabel pembelian detail berfungsi sebagai penyimpan rincian barang yang masuk pada setiap "
     "transaksi pembelian. Tabel ini menghubungkan data pembelian dengan produk spesifik melalui "
     "id_pembelian dan id_produk, serta merekam jumlah masuk (qty_masuk) dan harga modal "
     "(harga_masuk)."),

    # ---------------- Entitas Manajemen Layanan ----------------
    ("Entitas Manajemen Layanan (Garansi & Retur)", "5.12", "Tabel Klaim Garansi",
     "tb_klaim_garansi",
     [("id_klaim", "PK"), ("id_transaksi", "FK"), ("id_produk", "FK"),
      ("deskripsi_klaim", ""), ("biaya_klaim", ""), ("status_klaim", ""),
      ("tanggal_tambah_klaim", ""), ("tanggal_update_klaim", "")],
     "Tabel klaim garansi mencatat data permohonan garansi dari pelanggan atas produk yang telah "
     "dibeli. Relasi dengan id_transaksi dan id_produk memastikan bahwa barang yang diklaim memang "
     "valid dan terdaftar dalam riwayat transaksi penjualan, sehingga proses klaim dapat "
     "dipertanggungjawabkan."),

    ("Entitas Manajemen Layanan (Garansi & Retur)", "5.13", "Tabel Retur", "tb_retur",
     [("id_retur", "PK"), ("id_produk", "FK"), ("id_supplier", "FK"),
      ("tipe_retur", ""), ("qty_retur", ""), ("alasan_retur", ""),
      ("tanggal_tambah_retur", ""), ("tanggal_update_retur", "")],
     "Tabel retur digunakan untuk mencatat pengembalian barang, baik kepada supplier maupun dari "
     "pelanggan, akibat adanya cacat atau ketidaksesuaian barang. Tabel ini merekam alasan "
     "(alasan_retur), kuantitas (qty_retur), dan tipe retur (tipe_retur), serta memiliki relasi ke "
     "tabel produk dan supplier melalui kunci asing."),
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(dxa))
    tcPr.append(tcW)


def set_repeat_header(row):
    """Mark the row as a repeating header (cross-page)."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def add_run(paragraph, text, *, bold=False, size=None, color=None, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Calibri"
    return r


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()

# Default style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.3
normal.paragraph_format.space_after = Pt(6)

# Page margins (match original: ~2cm L/R, 1.6cm T/B)
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ----- Title -----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title, "Spesifikasi Tabel Database", bold=True, size=16, color="1F3864")
title.paragraph_format.space_after = Pt(8)

# ----- Intro -----
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(intro,
        "Spesifikasi tabel database bertujuan untuk memberikan gambaran teknis mengenai cara "
        "penyimpanan dan pengorganisasian data dalam sistem basis data. Berikut adalah tabel-tabel "
        "utama hasil proses normalisasi (3NF) yang terdapat dalam basis data sistem informasi "
        "Point of Sale (POS) UMKM El-Tech Pontianak:")

# Column widths (twips): No=720, Nama Atribut=3600, Keterangan=3600
COL_W = [720, 3600, 3600]
HEADER_FILL = "1F3864"   # dark blue
PK_FILL = "FCE4D6"       # light orange
FK_FILL = "DDEBF7"       # light blue

current_group = None
group_count = 0

for group, tbl_no, tbl_title, tbl_name, columns, desc in DATA:
    # ----- Group header (only when group changes) -----
    if group != current_group:
        current_group = group
        group_count += 1
        doc.add_paragraph()  # spacing
        gh = doc.add_paragraph()
        add_run(gh, group, bold=True, size=13, color="2E74B5")
        gh.paragraph_format.space_before = Pt(10)
        gh.paragraph_format.space_after = Pt(6)

    # ----- Table caption -----
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(2)
    add_run(cap, f"Tabel {tbl_no}  ", bold=True, size=11, color="1F3864")
    add_run(cap, tbl_title, bold=True, size=11, color="1F3864")
    add_run(cap, f"  ({tbl_name})", italic=True, size=10, color="808080")
    cap.paragraph_format.keep_with_next = True

    # ----- Build the table -----
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    set_repeat_header(hdr)
    headers = ["No", "Nama Atribut", "Keterangan (PK/FK)"]
    for i, htext in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, COL_W[i])
        set_cell_shading(cell, HEADER_FILL)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.name = "Calibri"

    # Data rows
    for idx, (attr, note) in enumerate(columns, start=1):
        row = table.add_row()
        # No
        c0 = row.cells[0]
        set_cell_width(c0, COL_W[0])
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = c0.paragraphs[0].add_run(str(idx))
        r0.font.size = Pt(10)
        r0.font.name = "Calibri"
        # Attribute
        c1 = row.cells[1]
        set_cell_width(c1, COL_W[1])
        # Bold the PK row's attribute
        is_pk = (note == "PK")
        r1 = c1.paragraphs[0].add_run(attr)
        r1.font.size = Pt(10)
        r1.font.name = "Consolas"   # monospace for column names
        r1.bold = is_pk
        # Keterangan
        c2 = row.cells[2]
        set_cell_width(c2, COL_W[2])
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if note:
            label = "Primary Key" if note == "PK" else "Foreign Key"
            r2 = c2.paragraphs[0].add_run(label)
            r2.font.size = Pt(10)
            r2.bold = True
            r2.font.color.rgb = RGBColor.from_string(
                "C00000" if note == "PK" else "1F4E79"
            )
            r2.font.name = "Calibri"
            set_cell_shading(c2, PK_FILL if note == "PK" else FK_FILL)
        else:
            r2 = c2.paragraphs[0].add_run("\u2013")  # en dash
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor.from_string("808080")

    # ----- Description paragraph -----
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    dp.paragraph_format.space_before = Pt(4)
    add_run(dp, desc, size=11)

# ----- Footer note -----
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(foot,
        "Catatan: PK = Primary Key,  FK = Foreign Key.",
        italic=True, size=9, color="808080")

# ----- Save -----
output = "Spesifikasi Tabel Database - 3NF Final.docx"
doc.save(output)
print(f"OK -> {output}")
print(f"Total tabel: {len(DATA)}")
