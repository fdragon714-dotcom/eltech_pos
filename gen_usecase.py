from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Set paragraph spacing to 0
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DESKRIPSI USE CASE\nSISTEM INFORMASI POS UMKM EL-TECH PONTIANAK')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

use_cases = [
    {
        'no': '5.18',
        'nama': 'Mengelola Klaim Garansi',
        'aktor': 'Admin, Kasir',
        'penjelasan': (
            'Use case ini mendeskripsikan aktivitas yang dilakukan oleh aktor Admin dan Kasir '
            'dalam mengelola data pengajuan klaim garansi dari pelanggan terhadap produk laptop '
            'yang telah dibeli. Admin memiliki kewenangan untuk melihat, menambah, mengubah, '
            'dan menghapus data klaim garansi, sedangkan Kasir hanya dapat melihat daftar klaim '
            'yang sudah tercatat. Setiap klaim garansi mencatat informasi transaksi penjualan '
            'terkait, produk yang diklaim, deskripsi kerusakan, biaya klaim, serta status klaim '
            'yang terdiri atas Pending, Diproses, atau Selesai. Proses ini penting dilaksanakan '
            'guna memantau dan mendokumentasikan seluruh klaim garansi yang diajukan oleh '
            'pelanggan secara tertib dan terpusat. Deskripsi lengkap alur use case Mengelola '
            'Klaim Garansi disajikan pada tabel berikut.'
        ),
        'pre': 'Pengguna telah login ke dalam sistem. Klaim garansi hanya dapat diajukan untuk transaksi penjualan yang sudah tercatat di database.',
        'flow': (
            '1. Admin atau Kasir memilih menu "Klaim Garansi" pada sidebar.\n'
            '2. Sistem menampilkan halaman daftar klaim garansi.\n'
            '3. Admin menekan tombol "+ Baru" untuk menambah klaim garansi baru.\n'
            '4. Admin memilih transaksi penjualan (invoice) yang menjadi acuan klaim.\n'
            '5. Sistem menampilkan daftar produk yang dibeli dalam transaksi tersebut.\n'
            '6. Admin memilih produk yang akan diklaim.\n'
            '7. Admin mengisi deskripsi kerusakan, biaya klaim, dan menetapkan status klaim.\n'
            '8. Admin menekan tombol "Save".\n'
            '9. Sistem menyimpan data klaim garansi ke database.\n'
            '10. Sistem menampilkan pesan bahwa data berhasil disimpan.'
        ),
        'post': 'Data klaim garansi berhasil tersimpan di database dan muncul pada daftar klaim garansi. Kasir hanya dapat melihat data klaim tanpa dapat menambah, mengubah, atau menghapus.',
    },
    {
        'no': '5.19',
        'nama': 'Mengelola Barang Masuk (Pembelian)',
        'aktor': 'Admin',
        'penjelasan': (
            'Use case ini mendeskripsikan aktivitas yang dilakukan oleh aktor Admin dalam '
            'mengelola data penerimaan barang masuk (pembelian) dari supplier ke dalam '
            'inventaris toko. Admin dapat melihat riwayat barang masuk, mencatat penerimaan '
            'barang masuk baru, serta menghapus data barang masuk yang sudah tercatat. '
            'Pada saat mencatat penerimaan barang baru, Admin memilih atau menambahkan '
            'supplier, memilih produk yang diterima, serta menentukan harga beli dan '
            'kuantitas setiap item. Sistem secara otomatis menghitung grand total pembelian, '
            'menambahkan stok produk, dan memperbarui Harga Pokok Penjualan (HPP) '
            'menggunakan metode rata-rata tertimbang (weighted average cost). Proses ini '
            'penting dilaksanakan guna memastikan seluruh barang yang diterima dari supplier '
            'tercatat dengan akurat dan nilai inventaris selalu mutakhir. Deskripsi lengkap '
            'alur use case Mengelola Barang Masuk disajikan pada tabel berikut.'
        ),
        'pre': 'Admin telah login ke dalam sistem. Terdapat produk yang terdaftar di database untuk diisi stoknya.',
        'flow': (
            '1. Admin memilih menu "Barang Masuk" pada sidebar.\n'
            '2. Sistem menampilkan halaman riwayat barang masuk.\n'
            '3. Admin menekan tombol "+ Baru" untuk mencatat penerimaan barang baru.\n'
            '4. Sistem menampilkan formulir penerimaan barang masuk.\n'
            '5. Admin memasukkan nama supplier (atau memilih dari daftar supplier yang sudah ada).\n'
            '6. Admin memilih produk yang diterima dari daftar produk.\n'
            '7. Admin memasukkan harga beli dan kuantitas setiap produk.\n'
            '8. Admin dapat menambahkan beberapa item produk sekaligus.\n'
            '9. Admin menekan tombol "Simpan".\n'
            '10. Sistem menyimpan data pembelian dan item-itemnya ke database.\n'
            '11. Sistem menambahkan stok setiap produk sesuai kuantitas yang dimasukkan.\n'
            '12. Sistem memperbarui Harga Pokok Penjualan (HPP) menggunakan metode weighted average cost.\n'
            '13. Sistem menampilkan pesan bahwa data berhasil disimpan.'
        ),
        'post': 'Data penerimaan barang masuk tersimpan di database. Stok produk bertambah dan HPP produk diperbarui. Admin juga dapat menghapus data barang masuk yang akan mengembalikan stok ke nilai sebelum penerimaan.',
    },
    {
        'no': '5.20',
        'nama': 'Melihat Stok Inventaris',
        'aktor': 'Admin, Kasir',
        'penjelasan': (
            'Use case ini mendeskripsikan aktivitas yang dilakukan oleh aktor Admin dan Kasir '
            'dalam melihat data inventaris stok produk secara lengkap. Halaman inventaris '
            'menampilkan seluruh produk yang terdaftar beserta kode SKU, nama produk, '
            'kategori, stok saat ini, Harga Pokok Penjualan (HPP), harga jual, total nilai '
            'aset berdasarkan HPP, total nilai aset berdasarkan harga jual, serta potensi '
            'laba dari seluruh stok yang tersedia. Kasir dapat mengakses halaman ini untuk '
            'memantau ketersediaan stok produk secara langsung, sedangkan Admin memiliki '
            'akses tambahan untuk melakukan penyesuaian stok secara manual. Proses ini '
            'penting dilaksanakan guna memberikan informasi stok yang akurat dan mendukung '
            'pengambilan keputusan dalam manajemen persediaan barang. Deskripsi lengkap '
            'alur use case Melihat Stok Inventaris disajikan pada tabel berikut.'
        ),
        'pre': 'Pengguna telah login ke dalam sistem dan terdapat produk yang terdaftar di database.',
        'flow': (
            '1. Pengguna memilih menu "Inventory" pada sidebar.\n'
            '2. Sistem memuat data seluruh produk dari database.\n'
            '3. Sistem menampilkan daftar produk dalam bentuk tabel yang berisi kode SKU, nama produk, kategori, stok saat ini, HPP, dan harga jual.\n'
            '4. Sistem menghitung dan menampilkan total nilai aset berdasarkan HPP (stok x HPP).\n'
            '5. Sistem menghitung dan menampilkan total nilai aset berdasarkan harga jual (stok x harga jual).\n'
            '6. Sistem menghitung dan menampilkan potensi laba (total nilai jual dikurangi total nilai HPP).'
        ),
        'post': 'Sistem menampilkan halaman inventaris dengan seluruh informasi stok dan ringkasan keuangan inventaris. Admin dapat menyesuaikan stok secara manual melalui tombol "Edit Stok", sedangkan Kasir hanya dapat melihat data inventaris.',
    },
    {
        'no': '5.21',
        'nama': 'Mengelola Retur Barang',
        'aktor': 'Admin',
        'penjelasan': (
            'Use case ini mendeskripsikan hak mutlak yang dimiliki oleh aktor Admin selaku '
            'pemilik (owner) sistem dalam mengelola data retur barang. Melalui fitur ini, '
            'Admin dapat mencatat retur barang baru, melihat riwayat retur, maupun menghapus '
            'data retur yang sudah tercatat. Retur barang mencakup dua tipe, yaitu '
            'pengembalian barang ke supplier dan pencatatan barang rusak atau cacat gudang. '
            'Sistem secara otomatis mengurangi stok produk sesuai kuantitas yang diretur. '
            'Deskripsi lengkap alur interaksi pengelolaan retur barang disajikan pada '
            'tabel berikut.'
        ),
        'pre': 'Admin telah login ke dalam sistem. Produk yang akan diretur tersedia di database dengan stok yang mencukupi.',
        'flow': (
            '1. Admin memilih menu "Retur Barang" pada sidebar.\n'
            '2. Sistem menampilkan halaman daftar retur barang.\n'
            '3. Admin menekan tombol "+ Baru" untuk mencatat retur baru.\n'
            '4. Sistem menampilkan formulir retur barang.\n'
            '5. Admin memilih produk yang akan diretur.\n'
            '6. Admin memilih tipe retur, yaitu "Kembalikan ke Supplier" atau "Barang Rusak Gudang".\n'
            '7. Apabila tipe retur adalah pengembalian ke supplier, Admin memilih supplier tujuan.\n'
            '8. Admin memasukkan kuantitas retur dan alasan pengembalian.\n'
            '9. Admin menekan tombol "Save".\n'
            '10. Sistem menyimpan data retur ke database.\n'
            '11. Sistem mengurangi stok produk sesuai kuantitas yang diretur.\n'
            '12. Sistem menampilkan pesan bahwa data berhasil disimpan.'
        ),
        'post': 'Data retur barang tersimpan di database dan stok produk berkurang sesuai kuantitas retur. Admin juga dapat menghapus data retur yang akan mengembalikan stok ke nilai sebelum retur.',
    }
]

for uc in use_cases:
    # Penjelasan paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(uc['penjelasan'])
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    
    # Create table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for cell in table.columns[0].cells:
        cell.width = Cm(4)
    for cell in table.columns[1].cells:
        cell.width = Cm(12)
    
    # Row 0: Merged header
    header_text = f'Tabel {uc["no"]} Deskripsi Use Case {uc["nama"]}'
    header_cell = table.rows[0].cells[0]
    header_cell_2 = table.rows[0].cells[1]
    
    tcPr = header_cell._tc.get_or_add_tcPr()
    merge = parse_xml(f'<w:tcMerge {nsdecls("w")} w:val="continue"/>')
    header_cell._tc.append(merge)
    
    tcPr2 = header_cell_2._tc.get_or_add_tcPr()
    merge2 = parse_xml(f'<w:tcMerge {nsdecls("w")} w:val="restart"/>')
    header_cell_2._tc.append(merge2)
    
    header_cell.text = header_text
    for paragraph in header_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    # Data rows
    rows_data = [
        ('Nama Use Case', uc['nama']),
        ('Aktor', uc['aktor']),
        ('Pre-kondisi', uc['pre']),
        ('Flow of Event', uc['flow']),
        ('Post-kondisi', uc['post']),
    ]
    
    for i, (label, value) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    # Styling
    for i in range(1, 6):
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        for paragraph in table.rows[i].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    
    doc.add_paragraph()

doc.save('Deskripsi_Use_Case_v2.docx')
print("OK: Deskirpsi_Use_Case_v2.docx created with penjelasan + tabel")
