from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

title = doc.add_heading('Spesifikasi Tabel Database', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Spesifikasi tabel database bertujuan untuk memberikan gambaran teknis mengenai '
    'bagaimana data disimpan dan diorganisir dalam sistem basis data. Berikut adalah '
    'tabel utama yang terdapat dalam basis data Sistem Informasi POS UMKM El-Tech Pontianak '
    'setelah melalui proses normalisasi hingga tahap Third Normal Form (3NF).'
)

def add_table(doc, table_num, title_text, desc_text, columns):
    doc.add_paragraph()
    p = doc.add_paragraph(f'Tabel 5.{table_num} {title_text}')
    p.runs[0].bold = True
    doc.add_paragraph(desc_text)
    headers = ['No', 'Nama Kolom', 'Tipe Data', 'Panjang', 'NULL / NOT NULL', 'Default']
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        for p in hdr.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'
    for idx, col in enumerate(columns):
        row = table.add_row()
        vals = [str(idx+1), col[0], col[1], col[2], col[3], col[4]]
        for i, v in enumerate(vals):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [Cm(1), Cm(4), Cm(2.2), Cm(1.5), Cm(2.5), Cm(4.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

# ============================================================
# DATA MASTER
# ============================================================

add_table(doc, 1, 'Tabel User (tb_user)',
    'Tabel user digunakan untuk menyimpan data yang berkaitan langsung dengan proses '
    'login dan otorisasi pengguna. Di dalam models.py, tabel ini merupakan bawaan '
    'dari django.contrib.auth.models.User.',
    [
        ['id_user', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['username', 'VARCHAR', '150', 'NOT NULL', '-'],
        ['email', 'VARCHAR', '254', 'NOT NULL', '-'],
        ['password', 'VARCHAR', '128', 'NOT NULL', '-'],
        ['last_login', 'DATETIME', '-', 'NULL', 'NULL'],
        ['is_superuser', 'BOOLEAN', '-', 'NOT NULL', 'FALSE'],
        ['is_staff', 'BOOLEAN', '-', 'NOT NULL', 'FALSE'],
        ['is_active', 'BOOLEAN', '-', 'NOT NULL', 'TRUE'],
        ['date_joined', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

add_table(doc, 2, 'Tabel Kategori (tb_kategori)',
    'Tabel kategori digunakan untuk menyimpan data merek laptop (brand) seperti '
    'Lenovo, ASUS, atau HP. Di dalam models.py, tabel ini direpresentasikan oleh '
    'kelas Category.',
    [
        ['id_kategori', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['nama_kategori', 'VARCHAR', '100', 'NOT NULL', '-'],
        ['deskripsi_kategori', 'TEXT', '-', 'NULL', 'NULL'],
        ['gambar_kategori', 'VARCHAR', '255', 'NULL', 'NULL'],
        ['status_kategori', 'INTEGER', '-', 'NOT NULL', '1'],
        ['tanggal_tambah_kategori', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_kategori', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

add_table(doc, 3, 'Tabel Tipe Merek (tb_tipe)',
    'Tabel tipe digunakan untuk menyimpan data tipe atau series dari suatu merek, '
    'misalnya ThinkPad dan IdeaPad untuk merek Lenovo. Di dalam models.py, tabel ini '
    'direpresentasikan oleh kelas CategoryType dengan ForeignKey ke Category.',
    [
        ['id_tipe', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['id_kategori', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_kategori'],
        ['nama_tipe', 'VARCHAR', '100', 'NOT NULL', '-'],
        ['gambar_tipe', 'VARCHAR', '255', 'NULL', 'NULL'],
    ])

add_table(doc, 4, 'Tabel Produk (tb_produk)',
    'Tabel produk digunakan untuk menyimpan data seluruh unit laptop yang tercatat '
    'dalam inventaris. Di dalam models.py, tabel ini direpresentasikan oleh kelas '
    'Products dengan ForeignKey ke Category. Produk tidak terhubung langsung ke tipe.',
    [
        ['id_produk', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['id_kategori', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_kategori'],
        ['kode_produk', 'VARCHAR', '100', 'NOT NULL', '-'],
        ['nama_produk', 'VARCHAR', '255', 'NOT NULL', '-'],
        ['deskripsi_produk', 'TEXT', '-', 'NULL', 'NULL'],
        ['harga_jual', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['hpp', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['stok_produk', 'INTEGER', '-', 'NOT NULL', '0'],
        ['garansi_hari', 'INTEGER', '-', 'NOT NULL', '0'],
        ['status_produk', 'INTEGER', '-', 'NOT NULL', '1'],
        ['gambar_produk', 'VARCHAR', '255', 'NULL', 'NULL'],
        ['tanggal_tambah_produk', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_produk', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

add_table(doc, 5, 'Tabel Supplier (tb_supplier)',
    'Tabel supplier digunakan untuk menyimpan data pemasok barang yang menyuplai unit '
    'laptop ke dalam inventaris toko.',
    [
        ['id_supplier', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['nama_supplier', 'VARCHAR', '255', 'NOT NULL', '-'],
        ['kontak', 'VARCHAR', '50', 'NOT NULL', '-'],
        ['alamat', 'TEXT', '-', 'NULL', 'NULL'],
        ['status_supplier', 'INTEGER', '-', 'NOT NULL', '1'],
        ['tanggal_tambah_supplier', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_supplier', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

# ============================================================
# TRANSAKSI KELUAR (PENJUALAN POS)
# ============================================================

add_table(doc, 6, 'Tabel Transaksi (tb_transaksi)',
    'Tabel transaksi digunakan untuk menyimpan data header transaksi penjualan. '
    'Di dalam models.py, tabel ini direpresentasikan oleh kelas Sales. Setiap transaksi '
    'terikat pada satu user (kasir) melalui ForeignKey.',
    [
        ['id_transaksi', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['kode_transaksi', 'VARCHAR', '100', 'NOT NULL', '-'],
        ['sub_total', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['grand_total', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['diskon_amount', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['diskon', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['tendered_amount', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['amount_change', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['payment_method', 'VARCHAR', '50', 'NOT NULL', 'Tunai'],
        ['customer_name', 'VARCHAR', '255', 'NULL', 'NULL'],
        ['customer_wa', 'VARCHAR', '50', 'NULL', 'NULL'],
        ['customer_address', 'TEXT', '-', 'NULL', 'NULL'],
        ['id_user', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_user'],
        ['tanggal_tambah_transaksi', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_transaksi', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

add_table(doc, 7, 'Tabel Detail Transaksi (tb_transaksi_detail)',
    'Tabel detail transaksi digunakan untuk menyimpan data item yang dibeli dalam '
    'setiap transaksi. Di dalam models.py, tabel ini direpresentasikan oleh kelas '
    'salesItems dengan ForeignKey ke Sales dan Products.',
    [
        ['id_item', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['id_transaksi', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_transaksi'],
        ['id_produk', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_produk'],
        ['harga_item', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['qty_beli', 'INTEGER', '-', 'NOT NULL', '0'],
        ['total_item', 'DECIMAL', '12,2', 'NOT NULL', '0'],
    ])

# ============================================================
# TRANSAKSI MASUK (INVENTORY / PEMBELIAN)
# ============================================================

add_table(doc, 8, 'Tabel Pembelian (tb_pembelian)',
    'Tabel pembelian digunakan untuk menyimpan data header transaksi pembelian barang '
    'dari supplier.',
    [
        ['id_pembelian', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['kode_pembelian', 'VARCHAR', '100', 'NOT NULL', '-'],
        ['id_supplier', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_supplier'],
        ['grand_total_pembelian', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['tanggal_tambah_pembelian', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_pembelian', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

add_table(doc, 9, 'Tabel Detail Pembelian (tb_pembelian_detail)',
    'Tabel detail pembelian digunakan untuk menyimpan data item yang dibeli dalam '
    'setiap transaksi pembelian.',
    [
        ['id_item_masuk', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['id_pembelian', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_pembelian'],
        ['id_produk', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_produk'],
        ['harga_masuk', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['qty_masuk', 'INTEGER', '-', 'NOT NULL', '0'],
        ['total_masuk', 'DECIMAL', '12,2', 'NOT NULL', '0'],
    ])

# ============================================================
# MANAJEMEN AFTER-SALES (GARANSI & RETUR)
# ============================================================

add_table(doc, 10, 'Tabel Klaim Garansi (tb_klaim_garansi)',
    'Tabel klaim garansi digunakan untuk mencatat pengajuan klaim garansi dari '
    'pelanggan terhadap produk yang telah dibeli.',
    [
        ['id_klaim', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['id_transaksi', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_transaksi'],
        ['id_produk', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_produk'],
        ['deskripsi_klaim', 'TEXT', '-', 'NULL', 'NULL'],
        ['biaya_klaim', 'DECIMAL', '12,2', 'NOT NULL', '0'],
        ['status_klaim', 'VARCHAR', '50', 'NOT NULL', 'Pending'],
        ['tanggal_tambah_klaim', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_klaim', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

add_table(doc, 11, 'Tabel Retur Barang (tb_retur)',
    'Tabel retur digunakan untuk mencatat proses pengembalian barang ke supplier. '
    'Kolom id_supplier dapat bernilai NULL apabila retur terjadi karena kerusakan '
    'internal gudang (bukan retur ke supplier).',
    [
        ['id_retur', 'INTEGER', '-', 'NOT NULL', 'Auto Increment (PK)'],
        ['id_produk', 'INTEGER', '-', 'NOT NULL', 'FK ke tb_produk'],
        ['id_supplier', 'INTEGER', '-', 'NULL', 'FK ke tb_supplier'],
        ['tipe_retur', 'VARCHAR', '50', 'NOT NULL', '-'],
        ['qty_retur', 'INTEGER', '-', 'NOT NULL', '0'],
        ['alasan_retur', 'TEXT', '-', 'NULL', 'NULL'],
        ['tanggal_tambah_retur', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
        ['tanggal_update_retur', 'DATETIME', '-', 'NOT NULL', 'CURRENT_TIMESTAMP'],
    ])

doc.save('Spesifikasi Tabel Database_v3.docx')
print('OK')
